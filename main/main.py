import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import date, datetime


def generate_rental_agreement(landlord_name, tenant_name, premises_address, late_fee, lease_start_date, lease_end_date, rent_per_month, security_deposit):
    document = Document()
    
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    
    # Helper function to set font properties for a paragraph
    def set_font_properties(paragraph, font_name, font_size):
        for run in paragraph.runs:
            font = run.font
            font.name = font_name
            font.size = font_size
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Title
    title = document.add_heading("ROOM RENTAL AGREEMENT", level=1)
    
    
    
    # Introduction 
    introduction_text = f"This Lease Agreement (this “Agreement”) is made on {today_date} by and between {landlord_name} (“Landlord”) and {tenant_name} (“Tenant”). Each Tenant is jointly and severally liable to Landlord for full payment of rent and performance in accordance with all other terms of this Agreement. Each Landlord and Tenant may be referred to individually as a “Party” and collectively as the “Parties.” "
    paragraph = document.add_paragraph(introduction_text)
    

    # Section 1: Premises
    document.add_heading("1. Premises", level=1)
    premises_text = f"The premises leased is a room, with a shared bathroom, located at {premises_address} (the “Premises”). Parking is not included with the Premises."
    paragraph = document.add_paragraph(premises_text)
    

    # Section 2: Agreement to Lease
    document.add_heading("2. Agreement to Lease", level=2)
    agreement_text = "Landlord agrees to lease to Tenant and Tenant agrees to lease from Landlord, according to the terms and conditions set forth herein, the Premises."
    agreement_paragraph = document.add_paragraph(agreement_text)
    set_font_properties(agreement_paragraph, 'Times New Roman', Pt(12))

    # Section 3: Term
    document.add_heading("3. Term", level=2)
    term_text = f"This Agreement will be for a term beginning on {lease_start_date} and ending on {lease_end_date} (the “Term”)."
    term_paragraph = document.add_paragraph(term_text)
    set_font_properties(term_paragraph, 'Times New Roman', Pt(12))

    # Section 4: Rent
    document.add_heading("4. Rent", level=2)
    rent_text = f"Tenant will pay Landlord a monthly rent of ${rent_per_month:.2f} for the Term. Rent will be payable in advance and due on the 1st day of each month during the Term. The first rent payment is payable to Landlord when Tenant signs this Agreement. Rent for any period during the Term which is for less than one month will be a pro rata portion of the monthly installment. Rent will be paid to Landlord at Landlord's address provided herein (or to such other places as directed by Landlord) by mail or in person by one of the following methods: Cash, Money order, Electronic transfer, and will be payable in U.S. Dollars."
    rent_paragraph = document.add_paragraph(rent_text)
    set_font_properties(rent_paragraph, 'Times New Roman', Pt(12))

    # Section 5: Late Fee
    document.add_heading("5. Late Fee", level=2)
    late_fee_text = f"Rent paid after the 1st day of each month will be deemed as late; and if rent is not paid within five (5) day(s) after such due date, Tenant agrees to pay a late charge of ${late_fee:.2f}."
    late_fee_paragraph = document.add_paragraph(late_fee_text)
    set_font_properties(late_fee_paragraph, 'Times New Roman', Pt(12))
    
    # Properties for standard paragraphs
    set_font_properties(paragraph, 'Times New Roman', Pt(12))

    # Save the document with a filename based on the tenant's name
    file_name = f"{tenant_name}_Rental_Agreement{now_date_time}.docx"
    file_path = os.path.join(output_folder, file_name)
    document.save(file_path)

    return file_path

# Create the output folder if it doesn't exist
output_folder = "Rental Agreements"
os.makedirs(output_folder, exist_ok=True)

# Variables
landlord_name = "Victoria Stuckey"
tenant_name = "Sagwa"
premises_address = "10828 Fulton Ave, Austin, TX 78754"
late_fee = 35
lease_start_date = "August 25, 2023"
lease_end_date = "August 26, 2023"
rent_per_month = 880
security_deposit = 880
#today_text_date = date.today().strftime("%B %d, %Y")
today = date.today()
today_date = today.strftime("%B %d, %Y")
now = datetime.now()
now_date_time = now.strftime("%m%d%y%H%M")


rental_agreement_file = generate_rental_agreement(landlord_name, tenant_name, premises_address, late_fee, lease_start_date, lease_end_date, rent_per_month, security_deposit)
print(f"Rental agreement generated and saved as '{rental_agreement_file}'.")
